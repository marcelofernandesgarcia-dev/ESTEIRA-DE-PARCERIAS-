import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- WORD DOCX GENERATOR ---
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_left_border(cell, color_hex="2670E8", size="36"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color_hex)
    tcBorders.append(left)
    for b in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{b}')
        node.set(qn('w:val'), 'none')
        tcBorders.append(node)
    tcPr.append(tcBorders)

def add_callout(doc, text_p_list, color_hex="2670E8", bg_hex="F0F6FF"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(5.8)
    
    set_cell_background(cell, bg_hex)
    set_cell_left_border(cell, color_hex, size="36")
    set_cell_margins(cell, top=150, bottom=150, left=250, right=200)
    
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(4)
    
    if isinstance(text_p_list[0], tuple):
        header_text, is_bold = text_p_list[0]
        run = p0.add_run(header_text)
        run.bold = is_bold
        run.font.name = 'Verdana'
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(15, 23, 42)
    else:
        run = p0.add_run(text_p_list[0])
        run.font.name = 'Verdana'
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(50, 50, 50)
        
    for p_text in text_p_list[1:]:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        if isinstance(p_text, tuple):
            t, b = p_text
            run = p.add_run(t)
            run.bold = b
        else:
            run = p.add_run(p_text)
        run.font.name = 'Verdana'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(60, 60, 60)
        
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(6)

# Database for 9 generalized chapters conforming to Gov.br guidelines
chapters_data = {
    "title": "CARTILHA DE NIVELAMENTO",
    "subtitle": "Módulo de Gestão de Parcerias e Ecossistema TRANSFEREGOV.BR\nGuia Prático para Órgãos Concedentes Federais, Estados, Municípios e OSCs",
    "meta_info": "Público: Concedentes, Estados, Municípios e OSCs",
    "expediente": [
        "Coordenação Geral: Diretoria de Parcerias e Transferências - DTPAR/SEGES/MGI",
        "Parceiros Institucionais: Estados, Municípios e Organizações da Sociedade Civil (OSCs)"
    ],
    "cap1": {
        "title": "Sumário Executivo",
        "icon": "⚡",
        "intro": "A modernização dos repasses de recursos no âmbito federal atinge um marco fundamental com a transição para o modelo de Transferência Simplificada no TRANSFEREGOV.BR. Esta mudança é uma resposta direta à necessidade de conferir celeridade e eficiência à execução de políticas públicas críticas de fomento e parcerias, otimizando o relacionamento interfederativo.",
        "points": [
            "• Ruptura de Paradigma: Dispensa convênios ou instrumentos complexos ex-ante, focando no objeto físico.",
            "• Padronização Técnica: Utiliza metas físicas pré-configuradas no Módulo de Gestão de Parcerias.",
            "• Mapeamento Estrutural: Uso de SIORG setorial (ex: Departamento Gestor - 267384) para controle administrativo.",
            "• Gestão de Riscos Eficiente: Regras de exceção (como o defeso eleitoral) parametrizadas por tipo de programa no sistema.",
            "• Rastreabilidade Financeira: Integração nativa com SIAFI para abertura de contas e conciliação em tempo real."
        ]
    },
    "cap2": {
        "title": "Contexto Institucional e Base Normativa",
        "icon": "🏛️",
        "p1": "A implementação desta nova modalidade é sustentada por Acordos de Cooperação Técnica (ACT) firmados entre o Ministério da Gestão e da Inovação em Serviços Públicos (MGI) e os órgãos federais parceiros. Este acordo formaliza o suporte técnico para a internalização de políticas finalísticas no ecossistema de transferências da União, substituindo a análise documental analítica por uma sistemática de adesão rápida.",
        "mgi": "Papel do MGI: Gestor central da Plataforma TRANSFEREGOV.BR, responsável por fornecer a infraestrutura sistêmica e garantir o marco regulatório e governança.",
        "parceiros": "Papel dos Órgãos Parceiros: Detêm total autonomia na gestão dos recursos de seus fundos, na definição das regras de negócio de suas chamadas públicas e na validação das metas físicas pactuadas."
    },
    "cap3": {
        "title": "Arquitetura Técnica do TRANSFEREGOV.BR",
        "icon": "⚙️",
        "intro": "A robustez do TRANSFEREGOV.BR garante a integridade dos recursos federais. A arquitetura sistêmica é subdividida em três ambientes principais:",
        "environments": [
            ("1. Treinamento", "Simulador (sandbox) isolado destinado ao aprendizado das equipes técnicas, sem impacto orçamentário real."),
            ("2. Homologação", "Validações finais de layouts, parametrizações e testes de integração de APIs com agentes bancários."),
            ("3. Produção", "Execução oficial das parcerias, com empenho no SIAFI, transações financeiras reais e monitoramento legal.")
        ],
        "siorg_solution": "Solução Técnica SIORG: Para viabilizar a vinculação administrativa de conselhos ou fundos finalísticos sem código de Unidade Gestora própria, a plataforma utiliza a vinculação via código de SIORG (ex: código do Departamento Gestor de Recursos - 267384) e perfis específicos para servidores."
    },
    "cap4": {
        "title": "Parametrização por Política Pública",
        "icon": "🔧",
        "intro": "O Módulo de Gestão de Parcerias do TRANSFEREGOV.BR se adapta às particularidades de diferentes políticas públicas por meio de configurações no sistema:",
        "items": [
            ("1. Requisitos de Habilitação", "Definição automatizada de critérios técnicos e fiscais que os proponentes devem cumprir eletronicamente."),
            ("2. Anexos Obrigatórios", "Possibilidade de restringir a adesão à inserção de planos setoriais específicos de trabalho (ex: planos de contingência, declarações locais)."),
            ("3. Metas Padronizadas", "Catálogo de entregas físicas predefinidas para a rápida adesão dos estados e municípios, eliminando propostas subjetivas.")
        ]
    },
    "cap5": {
        "title": "Ordem de Pagamento de Parcerias (OPP)",
        "icon": "💳",
        "intro": "A OPP é o instrumento financeiro definitivo de automatização bancária das parcerias. Ela assegura e agiliza a liberação dos recursos federais:",
        "pilares": [
            ("⚡ Tempo Real via API", "Conexão direta com bancos públicos oficiais, eliminando o processamento em lote (batch)."),
            ("🛡️ Dupla Assinatura", "Exigência de dupla autorização eletrônica (Operador e Gestor) com certificação Gov.br Ouro."),
            ("📱 Pix e QR Code", "Pagamento direto a fornecedores e prestadores na ponta, garantindo conciliação imediata."),
            ("⚙️ Controle Ativo", "Prerrogativa do órgão federal de comandar o bloqueio e devolução de saldos não utilizados.")
        ],
        "highlight": "Ficha Técnica e Prazos: Abertura de conta via API em 3 a 5 segundos; Janela de autorizações das 07h às 19h30; transações via Pix sem teto de valor. A migração compulsória para todas as 63 mil contas ativas no país deve ocorrer até dezembro de 2026.\n\nConsulta Pública: A consulta à OPP no Gestão de Parcerias está acessível na navegação em acesso livre através do endereço: https://parcerias.transferegov.sistema.gov.br/gestaofinanceira/opp/consulta . Exemplo de Parceria com OPP em produção (Piloto): 2024-00000623."
    },
    "cap6": {
        "title": "Jornada da Transferência",
        "icon": "🗺️",
        "intro": "A jornada de repasses simplificados substitui os planos de trabalho densos por um fluxo lógico em 5 etapas no sistema:",
        "steps": [
            ("01. Convocatória", "O órgão concedente cadastra a chamada pública, metas físicas padronizadas e prazos no sistema."),
            ("02. Proposição", "O ente subnacional adere digitalmente à chamada, preenchendo as informações e anexando os documentos."),
            ("03. Validação", "O concedente realiza a análise da habilitação eletrônica e emite o parecer de aprovação."),
            ("04. SIAFI", "O sistema realiza o empenho automático e o comando de abertura de conta bancária vinculada."),
            ("05. Ciclo Final", "O ente executa as metas físicas e apresenta o Relatório de Execução Simplificado focado nas entregas.")
        ]
    },
    "cap7": {
        "title": "Dúvidas Operacionais (FAQ)",
        "icon": "❓",
        "faqs": [
            ("Como funciona a adesão no TRANSFEREGOV.BR?", "O processo é eletrônico e integrado. O proponente localiza o edital e manifesta interesse aderindo diretamente no TRANSFEREGOV.BR, assinando digitalmente o termo."),
            ("O que significa parametrizar o edital?", "É a predefinição pelo órgão concedente das regras de negócio, metas padronizadas e documentos exigidos, adaptando o sistema ao seu setor."),
            ("Qual a importância da Sandbox de Treinamento?", "Funciona como um simulador (sandbox) isolado, permitindo testes práticos de empenho e execução de forma prática e sem risco fiscal real."),
            ("Como ocorre a prestação de contas?", "O foco é direcionado para a comprovação física do alcance do objeto e das metas cadastradas, simplificando a prestação documental ex-post.")
        ]
    },
    "cap8": {
        "title": "Pontos de Atenção e Riscos",
        "icon": "⚠️",
        "riscos": [
            ("Sintoma do Convênio (Resistência Cultural)", "Dificuldade cultural de técnicos em desapegar de controles ex-ante densos e Notas Fiscais unitárias ex-ante, o que invalida a celeridade proposta pela Transferência Simplificada."),
            ("Prazos do Defeso Eleitoral", "Observância estrita do calendário eleitoral ordinário. Importante notar que apenas programas de Emergência ou Calamidade Pública formalmente decretada por portaria ministerial ativa possuem exceções de continuidade.")
        ]
    },
    "cap9": {
        "title": "Boas Práticas Operacionais",
        "icon": "✅",
        "items": [
            ("Treinamento Prévio em Sandbox", "Utilização intensiva do ambiente de treinamento livre para simular 90% das telas sem necessidade imediata de certificado digital."),
            ("Diálogo Técnico Permanente", "Abertura de chats técnicos ou fóruns integrados para esclarecer parametrizações de metas com o órgão concedente antes de abrir o edital.")
        ]
    },
    "glossario": [
        ("TRANSFEREGOV.BR", "Plataforma integrada de operacionalização e controle das parcerias financeiras do Governo Federal."),
        ("SIORG", "Sistema de Informações Organizacionais do Governo Federal. Utilizado para mapeamento de responsabilidades administrativas no módulo."),
        ("SIAFI", "Sistema Integrado de Administração Financeira do Governo Federal. Responsável pelo empenho de fundos e abertura de contas bancárias vinculadas."),
        ("OPP", "Ordem de Pagamento de Parcerias. Canal de liquidação bancária automatizada e transparente em tempo real via Pix/QR Code."),
        ("ACT", "Acordo de Cooperação Técnica. Instrumento legal celebrado entre o MGI e o ministério parceiro para reger a utilização da Transferência Simplificada."),
        ("Concedente", "Órgão ou entidade da Administração Pública Federal responsável pela destinação do recurso e supervisão técnica."),
        ("Proponente", "Ente subnacional (Estado ou Município) ou OSC que manifesta interesse no edital e executa as metas pactuadas.")
    ]
}

def generate_docx():
    doc = docx.Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Cover Page
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run("Ministério da Gestão e da Inovação em Serviços Públicos\nSecretaria de Gestão e Inovação - SEGES")
    run_logo.font.name = 'Verdana'
    run_logo.font.size = Pt(11)
    run_logo.font.color.rgb = RGBColor(71, 85, 105)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(120)
    p_title.paragraph_format.space_after = Pt(8)
    run_title = p_title.add_run(chapters_data["title"])
    run_title.font.name = 'Verdana'
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(38, 112, 232)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(28)
    run_sub = p_sub.add_run(chapters_data["subtitle"])
    run_sub.font.name = 'Verdana'
    run_sub.font.size = Pt(12)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(71, 85, 105)
    
    p_year = doc.add_paragraph()
    p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_year.paragraph_format.space_before = Pt(140)
    run_year = p_year.add_run("Brasília - DF, 2026")
    run_year.font.name = 'Verdana'
    run_year.font.size = Pt(10)
    run_year.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_page_break()
    
    # Expediente
    h_exp = doc.add_paragraph()
    h_exp.paragraph_format.space_before = Pt(18)
    h_exp.paragraph_format.space_after = Pt(12)
    run_exp = h_exp.add_run("Expediente e Ficha Técnica")
    run_exp.font.name = 'Verdana'
    run_exp.bold = True
    run_exp.font.size = Pt(16)
    run_exp.font.color.rgb = RGBColor(15, 23, 42)
    
    for item in chapters_data["expediente"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        run.font.name = 'Verdana'
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(51, 65, 85)
        
    doc.add_page_break()

    # Helpers
    def add_h(text, icon, level):
        h = doc.add_paragraph()
        h.paragraph_format.keep_with_next = True
        h.paragraph_format.space_before = Pt(16 if level==1 else 12)
        h.paragraph_format.space_after = Pt(8 if level==1 else 6)
        full_text = f"{icon} {text}" if icon else text
        run = h.add_run(full_text)
        run.font.name = 'Verdana'
        run.bold = True
        run.font.size = Pt(14 if level==1 else 11.5)
        run.font.color.rgb = RGBColor(15, 23, 42) if level==1 else RGBColor(38, 112, 232)
        return h

    def add_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Verdana'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(51, 65, 85)
        return p

    # 1. Sumário Executivo
    add_h(chapters_data["cap1"]["title"], chapters_data["cap1"]["icon"], 1)
    add_p(chapters_data["cap1"]["intro"])
    add_callout(doc, [("Importante — Principais Conclusões:", True)] + chapters_data["cap1"]["points"], color_hex="2670E8", bg_hex="F0F6FF")

    # 2. Contexto Institucional
    add_h(chapters_data["cap2"]["title"], chapters_data["cap2"]["icon"], 1)
    add_p(chapters_data["cap2"]["p1"])
    add_callout(doc, [
        ("Mensagem Central:", True),
        chapters_data["cap2"]["mgi"],
        chapters_data["cap2"]["parceiros"]
    ], color_hex="3B82F6", bg_hex="F1F5F9")

    # 3. Arquitetura Técnica
    add_h(chapters_data["cap3"]["title"], chapters_data["cap3"]["icon"], 1)
    add_p(chapters_data["cap3"]["intro"])
    
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Shading Accent 1'
    for col_idx, text in enumerate(["Ambiente", "Descrição e Finalidade"]):
        cell = table.cell(0, col_idx)
        cell.text = text
        set_cell_background(cell, "2670E8")
        p = cell.paragraphs[0]
        run = p.runs[0]
        run.bold = True
        run.font.name = 'Verdana'
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    for idx, (env_name, env_desc) in enumerate(chapters_data["cap3"]["environments"]):
        row = table.rows[idx + 1]
        c0, c1 = row.cells[0], row.cells[1]
        c0.text = env_name
        set_cell_background(c0, "F1F5F9")
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.runs[0].bold = True
        p0.runs[0].font.name = 'Verdana'
        p0.runs[0].font.size = Pt(9.5)
        
        c1.text = env_desc
        p1 = c1.paragraphs[0]
        p1.runs[0].font.name = 'Verdana'
        p1.runs[0].font.size = Pt(9.5)
        
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(6)
    add_p(chapters_data["cap3"]["siorg_solution"])

    # 4. Parametrização
    add_h(chapters_data["cap4"]["title"], chapters_data["cap4"]["icon"], 1)
    add_p(chapters_data["cap4"]["intro"])
    for title, desc in chapters_data["cap4"]["items"]:
        add_h(title, None, 2)
        add_p(desc)

    # 5. OPP
    add_h(chapters_data["cap5"]["title"], chapters_data["cap5"]["icon"], 1)
    add_p(chapters_data["cap5"]["intro"])
    for title, desc in chapters_data["cap5"]["pilares"]:
        add_h(title, None, 2)
        add_p(desc)
    add_callout(doc, [("Importante — Ficha Técnica e Prazos:", True), chapters_data["cap5"]["highlight"]], color_hex="2670E8", bg_hex="F0F6FF")

    # 6. Jornada
    add_h(chapters_data["cap6"]["title"], chapters_data["cap6"]["icon"], 1)
    add_p(chapters_data["cap6"]["intro"])
    for title, desc in chapters_data["cap6"]["steps"]:
        add_p(f"• {title}: {desc}")

    # 7. FAQ
    add_h(chapters_data["cap7"]["title"], chapters_data["cap7"]["icon"], 1)
    for q, a in chapters_data["cap7"]["faqs"]:
        add_callout(doc, [(f"❓ {q}", True), a], color_hex="2670E8", bg_hex="F0F6FF")

    # 8. Riscos
    add_h(chapters_data["cap8"]["title"], chapters_data["cap8"]["icon"], 1)
    for title, desc in chapters_data["cap8"]["riscos"]:
        add_callout(doc, [(f"Atenção — {title}", True), desc], color_hex="EF4444", bg_hex="FEF2F2")

    # 9. Boas Práticas
    add_h(chapters_data["cap9"]["title"], chapters_data["cap9"]["icon"], 1)
    for title, desc in chapters_data["cap9"]["items"]:
        add_callout(doc, [(f"Mensagem Central — {title}", True), desc], color_hex="10B981", bg_hex="F0FDF4")

    # Glossário
    doc.add_page_break()
    add_h("Glossário Técnico", "📖", 1)
    for term, definition in chapters_data["glossario"]:
        add_callout(doc, [(term, True), definition], color_hex="2670E8", bg_hex="F8FAFC")

    # Contracapa
    doc.add_page_break()
    p_cc = doc.add_paragraph()
    p_cc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cc.paragraph_format.space_before = Pt(120)
    run_cc = p_cc.add_run("Secretaria de Gestão e Inovação (SEGES)\nMinistério da Gestão e da Inovação em Serviços Públicos\n\nCanais de Suporte:\nPortal Transferegov: www.gov.br/transferegov\nEmail: suporte.transferegov@gestao.gov.br")
    run_cc.font.name = 'Verdana'
    run_cc.font.size = Pt(10.5)
    run_cc.font.color.rgb = RGBColor(71, 85, 105)

    docx_path = r"C:\Users\marce\Documents\antigravity\charming-einstein\landing_transferegov\materiais_complementares\slides\Cartilha_de_Nivelamento_Modulo_de_Gestao_de_Parcerias_e_TRANSFEREGOV.docx"
    doc.save(docx_path)
    print("DOCX successfully generated with Verdana and Gov.br guidelines!")

# --- HTML TEMPLATE GENERATOR ---
def generate_html_files():
    html_template = """<!DOCTYPE html>
<html lang="pt-BR">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Cartilha de Nivelamento — Módulo de Gestão de Parcerias e TRANSFEREGOV</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Verdana&display=swap" rel="stylesheet">
  <style>
    :root {
      /* Primárias institucionais */
      --gov-blue-warm-vivid-60: hsl(215, 45%, 45%);
      --gov-blue-warm-vivid-10: hsl(210, 60%, 90%);
      --gov-indigo-cool-60: hsl(226, 45%, 45%);

      /* Feedback */
      --gov-green-cool-vivid-50: hsl(158, 100%, 25%);
      --gov-yellow-vivid-20: hsl(46, 100%, 74%);
      --gov-red-vivid-50: hsl(0, 100%, 38%);

      /* Neutros */
      --gov-gray-90: hsl(0, 0%, 11%);
      --gov-gray-60: hsl(0, 0%, 39%);
      --gov-gray-20: hsl(0, 0%, 80%);
      --gov-gray-10: hsl(0, 0%, 95%);
      --gov-pure-0: #FFFFFF;
      --gov-pure-100: #000000;

      --font-main: 'Verdana', sans-serif;
    }

    * {
      box-sizing: border-box;
      margin: 0;
      padding: 0;
    }

    body {
      font-family: var(--font-main);
      color: var(--gov-gray-90);
      background-color: var(--gov-gray-10);
      line-height: 1.6;
      font-size: 14px;
    }

    .document-page {
      max-width: 210mm;
      min-height: 297mm;
      margin: 30px auto;
      background: var(--gov-pure-0);
      padding: 22mm 20mm;
      border-radius: 8px;
      box-shadow: 0 4px 16px rgba(0,0,0,0.1);
      border: 1px solid var(--gov-gray-20);
      position: relative;
    }

    .action-bar {
      position: fixed;
      top: 20px;
      right: 20px;
      z-index: 1000;
      display: flex;
      gap: 10px;
    }

    .btn-action {
      background: var(--gov-blue-warm-vivid-60);
      color: var(--gov-pure-0);
      border: none;
      padding: 10px 18px;
      border-radius: 6px;
      font-weight: 600;
      font-size: 0.9rem;
      cursor: pointer;
      display: inline-flex;
      align-items: center;
      gap: 8px;
      box-shadow: 0 2px 6px rgba(0,0,0,0.15);
      transition: all 0.2s ease;
      text-decoration: none;
    }

    .btn-action:hover {
      background: var(--gov-indigo-cool-60);
      transform: translateY(-2px);
    }

    .cover-page {
      background: linear-gradient(135deg, var(--gov-blue-warm-vivid-60) 0%, var(--gov-indigo-cool-60) 100%);
      color: var(--gov-pure-0);
      padding: 50px 40px;
      border-radius: 12px;
      margin-bottom: 30px;
      border-left: 6px solid var(--gov-yellow-vivid-20);
      position: relative;
    }

    .cover-badge {
      display: inline-block;
      background: rgba(255, 255, 255, 0.15);
      color: var(--gov-yellow-vivid-20);
      border: 1px solid rgba(255, 255, 255, 0.3);
      padding: 4px 10px;
      border-radius: 4px;
      font-size: 0.75rem;
      font-weight: 700;
      letter-spacing: 1px;
      text-transform: uppercase;
      margin-bottom: 15px;
    }

    .cover-title {
      font-size: 24px;
      font-weight: 800;
      color: var(--gov-pure-0);
      line-height: 1.2;
      margin-bottom: 10px;
    }

    .cover-subtitle {
      font-size: 16px;
      color: var(--gov-gray-10);
      margin-bottom: 20px;
    }

    .cover-meta {
      font-size: 12px;
      color: var(--gov-gray-10);
      border-top: 1px solid rgba(255, 255, 255, 0.2);
      padding-top: 15px;
      display: flex;
      justify-content: space-between;
    }

    .expediente-section {
      background: var(--gov-gray-10);
      border-radius: 8px;
      padding: 24px;
      margin-bottom: 30px;
      border: 1px solid var(--gov-gray-20);
    }

    .expediente-section h3 {
      font-size: 18px;
      font-weight: 600;
      margin-bottom: 12px;
      color: var(--gov-blue-warm-vivid-60);
    }

    .expediente-list {
      list-style: none;
    }

    .expediente-list li {
      font-size: 14px;
      margin-bottom: 6px;
      color: var(--gov-gray-60);
    }

    .sumario-section {
      background: var(--gov-pure-0);
      border: 1px solid var(--gov-gray-20);
      border-radius: 8px;
      padding: 20px;
      margin-bottom: 30px;
    }

    .sumario-section h3 {
      font-size: 18px;
      font-weight: 600;
      margin-bottom: 12px;
      color: var(--gov-blue-warm-vivid-60);
    }

    .sumario-links {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 8px 24px;
      list-style-type: square;
      padding-left: 20px;
    }

    .sumario-links a {
      color: var(--gov-blue-warm-vivid-60);
      text-decoration: none;
      font-size: 14px;
    }

    .sumario-links a:hover {
      text-decoration: underline;
    }

    .section-container {
      margin-bottom: 40px;
      padding-top: 20px;
    }

    .num-header-wrapper {
      display: flex;
      align-items: flex-start;
      gap: 16px;
      margin-bottom: 14px;
    }

    .section-num {
      font-size: 48px;
      font-weight: 800;
      line-height: 1;
      color: var(--gov-blue-warm-vivid-60);
      font-family: var(--font-main);
    }

    h2.section-heading {
      font-size: 24px;
      font-weight: 700;
      color: var(--gov-gray-90);
      border-bottom: 2px solid var(--gov-gray-20);
      padding-bottom: 6px;
      flex: 1;
    }

    h3.sub-heading {
      font-size: 18px;
      font-weight: 600;
      color: var(--gov-blue-warm-vivid-60);
      margin-top: 16px;
      margin-bottom: 8px;
    }

    p {
      margin-bottom: 12px;
      text-align: justify;
      color: var(--gov-gray-60);
      font-size: 14px;
    }

    .grid-2 {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 16px;
      margin: 16px 0;
    }

    .grid-3 {
      display: grid;
      grid-template-columns: repeat(3, 1fr);
      gap: 16px;
      margin: 16px 0;
    }

    .card {
      background: var(--gov-gray-10);
      border: 1px solid var(--gov-gray-20);
      border-radius: 8px;
      padding: 16px;
    }

    .card h4 {
      font-size: 14px;
      color: var(--gov-gray-90);
      margin-bottom: 8px;
      font-weight: 700;
    }

    .card p {
      font-size: 14px;
      color: var(--gov-gray-60);
      margin: 0;
    }

    .callout {
      padding: 16px;
      border-radius: 8px;
      margin: 16px 0;
      font-size: 14px;
    }

    .callout-title {
      font-weight: 700;
      margin-bottom: 6px;
      color: var(--gov-gray-90);
    }

    /* Gov.br Specific Callouts */
    .callout-atencao { 
      border-left: 4px solid var(--gov-yellow-vivid-20); 
      background: var(--gov-yellow-vivid-20); 
      color: var(--gov-gray-90);
    }
    
    .callout-importante { 
      border-left: 4px solid var(--gov-blue-warm-vivid-60); 
      background: var(--gov-blue-warm-vivid-10); 
      color: var(--gov-gray-90);
    }
    
    .callout-mensagem { 
      border-left: 4px solid var(--gov-indigo-cool-60); 
      background: var(--gov-gray-10); 
      color: var(--gov-gray-90);
    }

    .callout-perigo {
      border-left: 4px solid var(--gov-red-vivid-50);
      background: rgba(239, 68, 68, 0.1);
      color: var(--gov-gray-90);
    }

    table.data-table {
      width: 100%;
      border-collapse: collapse;
      margin: 16px 0;
      font-size: 14px;
    }

    table.data-table th, 
    table.data-table td {
      border: 1px solid var(--gov-gray-20);
      padding: 10px 14px;
      text-align: left;
    }

    table.data-table th {
      background: var(--gov-blue-warm-vivid-60);
      font-weight: 700;
      color: var(--gov-pure-0);
    }

    table.data-table tr:nth-child(even) td {
      background: var(--gov-gray-10);
    }

    .badge {
      display: inline-block;
      padding: 2px 6px;
      border-radius: 4px;
      font-size: 12px;
      font-weight: 700;
      text-transform: uppercase;
    }

    .badge-primary { background: var(--gov-blue-warm-vivid-10); color: var(--gov-blue-warm-vivid-60); }
    .badge-success { background: rgba(16, 185, 129, 0.2); color: var(--gov-green-cool-vivid-50); }
    .badge-danger { background: rgba(239, 68, 68, 0.2); color: var(--gov-red-vivid-50); }

    .image-caption {
      font-size: 12px;
      color: var(--gov-gray-60);
      margin-top: 4px;
      text-align: center;
      font-style: italic;
    }

    .contracapa-section {
      margin-top: 60px;
      border-top: 1px solid var(--gov-gray-20);
      padding-top: 40px;
      text-align: center;
    }

    .qr-container {
      display: flex;
      justify-content: center;
      gap: 30px;
      margin-top: 20px;
    }

    .qr-box {
      border: 1px solid var(--gov-gray-20);
      padding: 10px;
      border-radius: 8px;
      background: var(--gov-pure-0);
      max-width: 170px;
    }

    .qr-box img {
      width: 150px;
      height: 150px;
    }

    .qr-label {
      font-size: 11px;
      font-weight: 600;
      margin-top: 6px;
      color: var(--gov-gray-60);
    }

    @media (max-width: 768px) {
      .document-page {
        width: 100%;
        margin: 0;
        padding: 20px 15px;
        border-radius: 0;
        border: none;
      }
      .grid-2, .grid-3 {
        grid-template-columns: 1fr;
      }
      .sumario-links {
        grid-template-columns: 1fr;
      }
      .qr-container {
        flex-direction: column;
        align-items: center;
      }
    }

    @media print {
      body {
        background: none;
        color: var(--gov-pure-100);
        font-size: 12pt;
      }

      .action-bar {
        display: none !important;
      }

      .document-page {
        margin: 0;
        padding: 0;
        max-width: 100%;
        box-shadow: none;
        border-radius: 0;
        border: none;
        background: none;
      }

      .cover-page {
        background: var(--gov-gray-10) !important;
        color: var(--gov-pure-100) !important;
        border: 1px solid var(--gov-gray-20);
        border-left: 6px solid var(--gov-blue-warm-vivid-60) !important;
        -webkit-print-color-adjust: exact;
        print-color-adjust: exact;
      }

      .cover-title {
        color: var(--gov-pure-100) !important;
      }

      .cover-subtitle {
        color: var(--gov-gray-60) !important;
      }

      .cover-badge {
        background: var(--gov-pure-0) !important;
        border: 1px solid var(--gov-gray-60) !important;
        color: var(--gov-blue-warm-vivid-60) !important;
      }

      .cover-meta {
        color: var(--gov-gray-60) !important;
        border-top: 1px solid var(--gov-gray-20) !important;
      }

      h2.section-heading {
        color: var(--gov-pure-100) !important;
        border-bottom: 2px solid var(--gov-gray-20) !important;
      }

      h3.sub-heading {
        color: var(--gov-blue-warm-vivid-60) !important;
      }

      p {
        color: var(--gov-gray-90) !important;
      }

      .card {
        background: var(--gov-gray-10) !important;
        border: 1px solid var(--gov-gray-20) !important;
      }

      .callout-atencao { border-left: 4px solid var(--gov-yellow-vivid-20) !important; background: var(--gov-gray-10) !important; color: var(--gov-pure-100) !important; }
      .callout-importante { border-left: 4px solid var(--gov-blue-warm-vivid-60) !important; background: var(--gov-gray-10) !important; color: var(--gov-pure-100) !important; }
      .callout-mensagem { border-left: 4px solid var(--gov-indigo-cool-60) !important; background: var(--gov-gray-10) !important; color: var(--gov-pure-100) !important; }
      .callout-perigo { border-left: 4px solid var(--gov-red-vivid-50) !important; background: var(--gov-gray-10) !important; color: var(--gov-pure-100) !important; }

      table.data-table th {
        background: var(--gov-gray-20) !important;
        color: var(--gov-pure-100) !important;
      }

      table.data-table td {
        color: var(--gov-pure-100) !important;
      }

      table.data-table tr:nth-child(even) td {
        background: var(--gov-gray-10) !important;
      }

      .page-break {
        page-break-before: always;
      }

      .no-break {
        page-break-inside: avoid;
      }

      @page {
        size: A4;
        margin: 15mm;
      }
    }
  </style>
</head>
<body>

  <!-- Barra de Ações Rápidas -->
  <div class="action-bar">
    <button onclick="window.print()" class="btn-action">
      <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polyline points="6 9 6 2 18 2 18 9"></polyline><path d="M6 18H4a2 2 0 0 1-2-2v-5a2 2 0 0 1 2-2h16a2 2 0 0 1 2 2v5a2 2 0 0 1-2 2h-2"></path><rect x="6" y="14" width="12" height="8"></rect></svg>
      Imprimir / Salvar em PDF
    </button>
  </div>

  <main class="document-page">

    <!-- CAPA DA CARTILHA -->
    <header class="cover-page">
      <span class="cover-badge">Guia Oficial de Capacitação</span>
      <h1 class="cover-title">Cartilha de Nivelamento</h1>
      <p class="cover-subtitle">Módulo de Gestão de Parcerias e Ecossistema TRANSFEREGOV.BR</p>
      <div class="cover-meta">
        <span><strong>DTPAR/SEGES/MGI</strong> — 2026</span>
      </div>
    </header>

    <!-- EXPEDIENTE -->
    <section class="expediente-section no-break">
      <h3>Expediente e Ficha Técnica</h3>
      <ul class="expediente-list">
        <li><strong>Coordenação Geral:</strong> Diretoria de Parcerias e Transferências - DTPAR/SEGES/MGI</li>
        <li><strong>Parceiros Institucionais:</strong> Estados, Municípios e Organizações da Sociedade Civil (OSCs)</li>
      </ul>
    </section>

    <!-- SUMÁRIO -->
    <section class="sumario-section no-break">
      <h3>Sumário de Navegação</h3>
      <ul class="sumario-links">
        <li><a href="#cap1">1. Sumário Executivo</a></li>
        <li><a href="#cap2">2. Contexto Institucional</a></li>
        <li><a href="#cap3">3. Arquitetura Técnica</a></li>
        <li><a href="#cap4">4. Parametrização</a></li>
        <li><a href="#cap5">5. Ordem de Pagamento (OPP)</a></li>
        <li><a href="#cap6">6. Jornada da Transferência</a></li>
        <li><a href="#cap7">7. Dúvidas Frequentes (FAQ)</a></li>
        <li><a href="#cap8">8. Pontos de Atenção e Riscos</a></li>
        <li><a href="#cap9">9. Boas Práticas Operacionais</a></li>
        <li><a href="#glossario">Glossário de Termos</a></li>
      </ul>
    </section>

    <div class="page-break"></div>

    <!-- 1. SUMÁRIO EXECUTIVO -->
    <section class="section-container" id="cap1">
      <div class="num-header-wrapper">
        <span class="section-num" aria-hidden="true">01</span>
        <h2 class="section-heading">Sumário Executivo</h2>
      </div>
      <p>A modernização dos repasses de recursos no âmbito federal atinge um marco fundamental com a transição para o modelo de <strong>Transferência Simplificada</strong> no <strong>TRANSFEREGOV.BR</strong>. Esta mudança é uma resposta direta à necessidade de conferir celeridade e eficiência à execução de políticas públicas críticas de fomento e parcerias, otimizando o relacionamento interfederativo.</p>
      
      <div class="callout callout-importante">
        <div class="callout-title">Importante — Principais Conclusões:</div>
        <p>• <strong>Ruptura de Paradigma:</strong> Dispensa convênios ou instrumentos complexos ex-ante, focando no objeto físico.</p>
        <p>• <strong>Padronização Técnica:</strong> Utiliza metas físicas pré-configuradas no Módulo de Gestão de Parcerias.</p>
        <p>• <strong>Mapeamento Estrutural:</strong> Uso de SIORG setorial (ex: Departamento Gestor - 267384) para controle administrativo.</p>
        <p>• <strong>Gestão de Riscos Eficiente:</strong> Regras de exceção (como o defeso eleitoral) parametrizadas por tipo de programa no sistema.</p>
        <p>• <strong>Rastreabilidade Financeira:</strong> Integração nativa com SIAFI para abertura de contas e conciliação em tempo real.</p>
      </div>

      <div style="margin: 20px 0; text-align: center;">
        <img src="materiais_complementares/slides/Guia_das_Transferências_Simplificadas_Transferegov.png" alt="Guia das Transferências Simplificadas" style="max-width: 100%; height: auto; border-radius: 8px; border: 1px solid var(--gov-gray-20);">
        <p class="image-caption">Imagem ilustrativa — Visão geral da transição das parcerias no TRANSFEREGOV.BR.</p>
      </div>
    </section>

    <!-- 2. CONTEXTO INSTITUCIONAL -->
    <section class="section-container no-break" id="cap2">
      <div class="num-header-wrapper">
        <span class="section-num" aria-hidden="true">02</span>
        <h2 class="section-heading">Contexto Institucional e Base Normativa</h2>
      </div>
      <p>A implementação desta nova modalidade é sustentada por Acordos de Cooperação Técnica (ACT) firmados entre o Ministério da Gestão e da Inovação em Serviços Públicos (MGI) e os órgãos federais parceiros. Este acordo formaliza o suporte técnico para a internalização de políticas finalísticas no ecossistema de transferências da União, substituindo a análise documental analítica por uma sistemática de adesão rápida.</p>
      
      <div class="grid-2">
        <div class="card">
          <h4>Papel do MGI</h4>
          <p>Gestor central da Plataforma TRANSFEREGOV.BR, responsável por fornecer a infraestrutura sistêmica e garantir o marco regulatório e governança.</p>
        </div>
        <div class="card">
          <h4>Papel dos Órgãos Parceiros</h4>
          <p>Detêm total autonomia na gestão dos recursos de seus fundos, na definição das regras de negócio de suas chamadas públicas e na validação das metas físicas pactuadas.</p>
        </div>
      </div>
    </section>

    <!-- 3. ARQUITETURA TÉCNICA -->
    <section class="section-container no-break" id="cap3">
      <div class="num-header-wrapper">
        <span class="section-num" aria-hidden="true">03</span>
        <h2 class="section-heading">Arquitetura Técnica do TRANSFEREGOV.BR</h2>
      </div>
      <p>A robustez do TRANSFEREGOV.BR garante a integridade dos recursos federais. A arquitetura sistêmica é subdividida em três ambientes principais:</p>
      
      <div class="grid-3">
        <div class="card">
          <span class="badge badge-primary">1. Treinamento</span>
          <p>Simulador (sandbox) isolado destinado ao aprendizado das equipes técnicas, sem impacto orçamentário real.</p>
        </div>
        <div class="card">
          <span class="badge badge-primary">2. Homologação</span>
          <p>Validações finais de layouts, parametrizações e testes de integração de APIs com agentes bancários.</p>
        </div>
        <div class="card">
          <span class="badge badge-danger">3. Produção</span>
          <p>Execução oficial das parcerias, com empenho no SIAFI, transações financeiras reais e monitoramento legal.</p>
        </div>
      </div>
      
      <div class="callout callout-mensagem" style="margin-top:16px;">
        <div class="callout-title">Mensagem central:</div>
        <strong>Solução Técnica SIORG:</strong> Para viabilizar a vinculação administrativa de conselhos ou fundos finalísticos sem código de Unidade Gestora própria, a plataforma utiliza a vinculação via código de SIORG (ex: código do Departamento Gestor de Recursos - 267384) e perfis específicos para servidores.
      </div>
    </section>

    <!-- 4. PARAMETRIZAÇÃO -->
    <section class="section-container no-break" id="cap4">
      <div class="num-header-wrapper">
        <span class="section-num" aria-hidden="true">04</span>
        <h2 class="section-heading">Parametrização por Política Pública</h2>
      </div>
      <p>O Módulo de Gestão de Parcerias do TRANSFEREGOV.BR se adapta às particularidades de diferentes políticas públicas por meio de configurações no sistema:</p>
      
      <div class="card" style="margin-bottom:16px;">
        <strong>1. Requisitos de Habilitação:</strong> Definição automatizada de critérios técnicos e fiscais que os proponentes devem cumprir eletronicamente.
      </div>
      <div class="card" style="margin-bottom:16px;">
        <strong>2. Anexos Obrigatórios:</strong> Possibilidade de restringir a adesão à inserção de planos setoriais específicos de trabalho (ex: planos de contingência, declarações locais).
      </div>
      <div class="card">
        <strong>3. Metas Padronizadas:</strong> Catálogo de entregas físicas predefinidas para a rápida adesão dos estados e municípios, eliminando propostas subjetivas.
      </div>
    </section>

    <div class="page-break"></div>

    <!-- 5. OPP -->
    <section class="section-container" id="cap5">
      <div class="num-header-wrapper">
        <span class="section-num" aria-hidden="true">05</span>
        <h2 class="section-heading">Ordem de Pagamento de Parcerias (OPP)</h2>
      </div>
      <p>A OPP é o instrumento financeiro definitivo de automatização bancária das parcerias. Ela assegura e agiliza a liberação dos recursos federais:</p>
      
      <div class="grid-2">
        <div class="card">
          <h4>⚡ Tempo Real via API</h4>
          <p>Conexão direta com bancos públicos oficiais, eliminando o processamento em lote (batch).</p>
        </div>
        <div class="card">
          <h4>🛡️ Dupla Assinatura</h4>
          <p>Exigência de dupla autorização eletrônica (Operador e Gestor) com certificação Gov.br Ouro.</p>
        </div>
        <div class="card">
          <h4>📱 Pix e QR Code</h4>
          <p>Pagamento direto a fornecedores e prestadores na ponta, garantindo conciliação imediata.</p>
        </div>
        <div class="card">
          <h4>⚙️ Controle Ativo</h4>
          <p>Prerrogativa do órgão federal de comandar o bloqueio e devolução de saldos não utilizados.</p>
        </div>
      </div>
      
      <div class="callout callout-importante" style="margin-top:16px;">
        <div class="callout-title">Importante — Ficha Técnica e Prazos:</div>
        <p>Abertura de conta via API em 3 a 5 segundos; Janela de autorizações das 07h às 19h30; transações via Pix sem teto de valor. A migração compulsória para todas as 63 mil contas ativas no país deve ocorrer até dezembro de 2026.</p>
        <p style="margin-top: 10px;"><strong>Consulta Pública:</strong> A consulta à OPP no Gestão de Parcerias está acessível na navegação em acesso livre através do endereço: <a href="https://parcerias.transferegov.sistema.gov.br/gestaofinanceira/opp/consulta" target="_blank" rel="noopener noreferrer" style="color: var(--gov-blue-warm-vivid-60); text-decoration: underline;">https://parcerias.transferegov.sistema.gov.br/gestaofinanceira/opp/consulta</a>. Exemplo de Parceria com OPP em produção (Piloto): <strong>2024-00000623</strong>.</p>
      </div>

      <div style="margin: 20px 0; text-align: center;">
        <img src="materiais_complementares/OPP TREINAMENTO/Fluxo_da_Ordem_de_Pagamento.png" alt="Fluxo da Ordem de Pagamento" style="max-width: 100%; height: auto; border-radius: 8px; border: 1px solid var(--gov-gray-20);">
        <p class="image-caption">Imagem ilustrativa — Diagrama operacional da Ordem de Pagamento de Parcerias (OPP).</p>
      </div>
    </section>

    <!-- 6. JORNADA -->
    <section class="section-container no-break" id="cap6">
      <div class="num-header-wrapper">
        <span class="section-num" aria-hidden="true">06</span>
        <h2 class="section-heading">Jornada da Transferência</h2>
      </div>
      <p>A jornada de repasses simplificados substitui os planos de trabalho densos por um fluxo lógico em 5 etapas no sistema:</p>
      
      <table class="data-table">
        <thead>
          <tr>
            <th style="width:10%;">Etapa</th>
            <th style="width:25%;">Fase</th>
            <th>Descrição e Ação do Sistema</th>
          </tr>
        </thead>
        <tbody>
          <tr>
            <td><strong>01</strong></td>
            <td>Convocatória</td>
            <td>O órgão concedente cadastra a chamada pública, metas físicas padronizadas e prazos no sistema.</td>
          </tr>
          <tr>
            <td><strong>02</strong></td>
            <td>Manifestação de Interesse</td>
            <td>O ente subnacional adere digitalmente à chamada, preenchendo as informações e anexando os documentos.</td>
          </tr>
          <tr>
            <td><strong>03</strong></td>
            <td>Validação Técnica</td>
            <td>O concedente realiza a análise da habilitação eletrônica e emite o parecer de aprovação.</td>
          </tr>
          <tr>
            <td><strong>04</strong></td>
            <td>Empenho e Desembolso</td>
            <td>O sistema realiza o empenho automático e o comando de abertura de conta bancária vinculada.</td>
          </tr>
          <tr>
            <td><strong>05</strong></td>
            <td>Execução e Monitoramento</td>
            <td>O ente executa as metas físicas e apresenta o Relatório de Execução Simplificado focado nas entregas.</td>
          </tr>
        </tbody>
      </table>

      <div style="margin: 20px 0; text-align: center;">
        <img src="materiais_complementares/slides/Jornada_da_Transferência_Simplificada (1).png" alt="Jornada da Transferência Simplificada" style="max-width: 100%; height: auto; border-radius: 8px; border: 1px solid var(--gov-gray-20);">
        <p class="image-caption">Imagem ilustrativa — Trilha operacional interfederativa da proposição à entrega do objeto.</p>
      </div>
    </section>

    <!-- 7. FAQ -->
    <section class="section-container no-break" id="cap7">
      <div class="num-header-wrapper">
        <span class="section-num" aria-hidden="true">07</span>
        <h2 class="section-heading">Dúvidas Operacionais (FAQ)</h2>
      </div>
      
      <div class="callout callout-importante">
        <strong>Como funciona a adesão no TRANSFEREGOV.BR?</strong><br>
        O processo é eletrônico e integrado. O proponente localiza o edital e manifesta interesse aderindo diretamente no TRANSFEREGOV.BR, assinando digitalmente o termo.
      </div>
      <div class="callout callout-importante">
        <strong>O que significa parametrizar o edital?</strong><br>
        É a predefinição pelo órgão concedente das regras de negócio, metas padronizadas e documentos exigidos, adaptando o sistema ao seu setor.
      </div>
      <div class="callout callout-importante">
        <strong>Qual a importância da Sandbox de Treinamento?</strong><br>
        Funciona como um simulador (sandbox) isolado, permitindo testes práticos de empenho e execução de forma prática e sem risco fiscal real.
      </div>
      <div class="callout callout-importante">
        <strong>Como ocorre a prestação de contas?</strong><br>
        O foco é direcionado para a comprovação física do alcance do objeto e das metas cadastradas, simplificando a prestação documental ex-post.
      </div>
    </section>

    <!-- 8. RISCOS -->
    <section class="section-container no-break" id="cap8">
      <div class="num-header-wrapper">
        <span class="section-num" aria-hidden="true">08</span>
        <h2 class="section-heading">Pontos de Atenção e Riscos</h2>
      </div>
      
      <div class="callout callout-perigo">
        <div class="callout-title">Atenção — Sintoma do Convênio (Resistência Cultural):</div>
        Dificuldade cultural de técnicos em desapegar de controles ex-ante densos e Notas Fiscais unitárias ex-ante, o que invalida a celeridade proposta pela Transferência Simplificada.
      </div>
      
      <div class="callout callout-atencao">
        <div class="callout-title">Atenção — Prazos do Defeso Eleitoral:</div>
        Observância estrita do calendário eleitoral ordinário. Importante notar que apenas programas de Emergência ou Calamidade Pública formalmente decretada por portaria ministerial ativa possuem exceções de continuidade.
      </div>
    </section>

    <!-- 9. BOAS PRÁTICAS -->
    <section class="section-container no-break" id="cap9">
      <div class="num-header-wrapper">
        <span class="section-num" aria-hidden="true">09</span>
        <h2 class="section-heading">Boas Práticas Operacionais</h2>
      </div>
      
      <div class="callout callout-mensagem">
        <div class="callout-title">Mensagem central — Treinamento Prévio em Sandbox:</div>
        Utilização intensiva do ambiente de treinamento livre para simular 90% das telas sem necessidade imediata de certificado digital.
      </div>
      <div class="callout callout-mensagem">
        <div class="callout-title">Mensagem central — Diálogo Técnico Permanente:</div>
        Abertura de chats técnicos ou fóruns integrados para esclarecer parametrizações de metas com o órgão concedente antes de abrir o edital.
      </div>

      <div style="margin: 20px 0; text-align: center;">
        <img src="materiais_complementares/slides/NotebookLM Mind Map (14).png" alt="Miniatura do Mapa Mental do Treinamento" style="max-width: 100%; height: auto; border-radius: 8px; border: 1px solid var(--gov-gray-20);">
        <p class="image-caption">Imagem ilustrativa — Mapa mental sintético de suporte à implantação e boas práticas.</p>
      </div>
    </section>

    <!-- GLOSSÁRIO -->
    <section class="section-container no-break" id="glossario">
      <h2 class="section-heading">📖 Glossário de Termos</h2>
      <p style="margin-top:10px;">Principais conceitos técnicos e siglas referenciados ao longo deste guia prático:</p>
      
      <div style="display: flex; flex-direction: column; gap: 12px; margin-top:14px;">
        <div class="card">
          <strong>TRANSFEREGOV.BR:</strong> Plataforma integrada de operacionalização e controle das parcerias financeiras do Governo Federal.
        </div>
        <div class="card">
          <strong>SIORG:</strong> Sistema de Informações Organizacionais do Governo Federal. Utilizado para mapeamento de responsabilidades administrativas no módulo.
        </div>
        <div class="card">
          <strong>SIAFI:</strong> Sistema Integrado de Administração Financeira do Governo Federal. Responsável pelo empenho de fundos e abertura de contas bancárias vinculadas.
        </div>
        <div class="card">
          <strong>OPP:</strong> Ordem de Pagamento de Parcerias. Canal de liquidação bancária automatizada e transparente em tempo real via Pix/QR Code.
        </div>
        <div class="card">
          <strong>ACT:</strong> Acordo de Cooperação Técnica. Instrumento legal celebrado entre o MGI e o ministério parceiro para reger a utilização da Transferência Simplificada.
        </div>
        <div class="card">
          <strong>Concedente:</strong> Órgão ou entidade da Administração Pública Federal responsável pela destinação do recurso e supervisão técnica.
        </div>
        <div class="card">
          <strong>Proponente:</strong> Ente subnacional (Estado ou Município) ou OSC que manifesta interesse no edital e executa as metas pactuadas.
        </div>
      </div>
    </section>

    <!-- CONTRACAPA -->
    <footer class="contracapa-section no-break">
      <p><strong>Secretaria de Gestão e Inovação (SEGES)</strong></p>
      <p>Ministério da Gestão e da Inovação em Serviços Públicos (MGI)</p>
      <p style="margin-top: 15px; font-size: 12px; color: var(--gov-gray-60);">Para mais informações ou suporte, acesse os portais oficiais:</p>
      
      <div class="qr-container">
        <div class="qr-box">
          <img src="https://api.qrserver.com/v1/create-qr-code/?size=150x150&data=https://tre-siconv.estaleiro.serpro.gov.br/ep-atos-prep-web/" alt="QR Code Sandbox Treinamento">
          <div class="qr-label">1. Sandbox Treinamento</div>
        </div>
        <div class="qr-box">
          <img src="https://api.qrserver.com/v1/create-qr-code/?size=150x150&data=https://www.gov.br/transferegov/" alt="QR Code Portal Transferegov">
          <div class="qr-label">2. Portal Transferegov</div>
        </div>
      </div>
    </footer>

  </main>
  
  [AUTO_PRINT_SCRIPT]

</body>
</html>"""

    html_normal = html_template.replace("[AUTO_PRINT_SCRIPT]", "")
    print_script = """<script>
    window.addEventListener('DOMContentLoaded', () => {
      setTimeout(() => {
        window.print();
      }, 350);
    });
  </script>"""
    html_download = html_template.replace("[AUTO_PRINT_SCRIPT]", print_script)

    dirs = [
        r"C:\Users\marce\Documents\antigravity\charming-einstein\landing_transferegov",
        r"C:\Users\marce\Documents\antigravity\charming-einstein\landing_treinamento_transferegov"
    ]
    
    for base in dirs:
        with open(os.path.join(base, "cartilha.html"), "w", encoding="utf-8") as f:
            f.write(html_normal)
        with open(os.path.join(base, "cartilha_download.html"), "w", encoding="utf-8") as f:
            f.write(html_download)
            
    print("HTML cartilha files compiled with Verdana, Gov.br palettes and QR codes successfully!")

# --- PATCH INDEX.HTML READER MODAL AND EXPEDIENTE FOOTER ---
def patch_index_modal():
    index_path = r"C:\Users\marce\Documents\antigravity\charming-einstein\landing_transferegov\index.html"
    if not os.path.exists(index_path):
        return
        
    with open(index_path, "r", encoding="utf-8") as f:
        html_content = f.read()

    # Define the new chapters block to be injected inside the modal-body in index.html (conforming to the updated structured content)
    new_chapters_html = """          <!-- Cap 1 -->
          <div class="report-chapter active">
            <button class="report-chapter-header" type="button" aria-expanded="true">⚡ 1. Sumário Executivo <span>▾</span></button>
            <div class="report-chapter-content" style="display: block;">
              <div class="cartilha-grid" style="display: grid; grid-template-columns: 1.2fr 0.8fr; gap: 2rem; align-items: center; margin-top: 1rem;">
                <div>
                  <p style="font-size: 0.95rem; line-height: 1.6; margin-bottom: 1.5rem; color: var(--text-secondary);">A modernização dos repasses de recursos no âmbito federal atinge um marco fundamental com a transição para o modelo de <strong>Transferência Simplificada</strong> no <strong>TRANSFEREGOV.BR</strong>. Esta mudança é uma resposta direta à necessidade de conferir celeridade e eficiência à execução de políticas públicas críticas de fomento e parcerias, otimizando o relacionamento interfederativo.</p>
                  <p style="font-size: 1rem; color: #fff; font-weight: 700; margin-bottom: 1rem; font-family: var(--font-display);">Principais Conclusões:</p>
                  <div style="display: flex; flex-direction: column; gap: 0.75rem;">
                    <div style="background: rgba(255,255,255,0.02); padding: 0.75rem 1rem; border-radius: 6px; border-left: 3px solid var(--secondary); font-size: 0.85rem;">
                      <strong>Ruptura de Paradigma:</strong> Dispensa convênios ou instrumentos complexos ex-ante, focando no objeto físico.
                    </div>
                    <div style="background: rgba(255,255,255,0.02); padding: 0.75rem 1rem; border-radius: 6px; border-left: 3px solid var(--secondary); font-size: 0.85rem;">
                      <strong>Padronização Técnica:</strong> Utiliza metas físicas pré-configuradas no Módulo de Gestão de Parcerias.
                    </div>
                    <div style="background: rgba(255,255,255,0.02); padding: 0.75rem 1rem; border-radius: 6px; border-left: 3px solid var(--secondary); font-size: 0.85rem;">
                      <strong>Mapeamento Estrutural:</strong> Uso de SIORG setorial (ex: Departamento Gestor - 267384) para controle administrativo.
                    </div>
                    <div style="background: rgba(255,255,255,0.02); padding: 0.75rem 1rem; border-radius: 6px; border-left: 3px solid var(--secondary); font-size: 0.85rem;">
                      <strong>Gestão de Riscos Eficiente:</strong> Regras de exceção (como o defeso eleitoral) parametrizadas por tipo de programa no sistema.
                    </div>
                    <div style="background: rgba(255,255,255,0.02); padding: 0.75rem 1rem; border-radius: 6px; border-left: 3px solid var(--secondary); font-size: 0.85rem;">
                      <strong>Rastreabilidade Financeira:</strong> Integração nativa com SIAFI para abertura de contas e conciliação em tempo real.
                    </div>
                  </div>
                </div>
                <div style="text-align: center; background: rgba(0,0,0,0.2); padding: 1rem; border-radius: 8px; border: 1px solid var(--border-color);">
                  <span style="font-size: 0.8rem; color: var(--text-muted); display: block; margin-bottom: 0.5rem; font-weight: bold; text-transform: uppercase; letter-spacing: 1px;">Visão Geral do Módulo</span>
                  <a href="materiais_complementares/slides/Guia_das_Transferências_Simplificadas_Transferegov.png" target="_blank" rel="noopener noreferrer">
                    <img src="materiais_complementares/slides/Guia_das_Transferências_Simplificadas_Transferegov.png" alt="Guia das Transferências Simplificadas" style="width: 100%; border-radius: 4px; box-shadow: 0 4px 12px rgba(0,0,0,0.4); margin-bottom: 0.5rem; transition: transform 0.2s;" onmouseover="this.style.transform='scale(1.02)'" onmouseout="this.style.transform='scale(1)'">
                  </a>
                  <span style="font-size: 0.75rem; color: var(--text-muted); display: block;">Clique para ampliar o mapa mental</span>
                </div>
              </div>
            </div>
          </div>

          <!-- Cap 2 -->
          <div class="report-chapter">
            <button class="report-chapter-header" type="button" aria-expanded="false">🏛️ 2. Contexto Institucional e Base Normativa <span>▸</span></button>
            <div class="report-chapter-content">
              <p style="font-size: 0.95rem; line-height: 1.6; margin-bottom: 1.5rem; color: var(--text-secondary);">A implementação desta nova modalidade é sustentada por Acordos de Cooperação Técnica (ACT) firmados entre o Ministério da Gestão e da Inovação em Serviços Públicos (MGI) e os órgãos federais parceiros. Este acordo formaliza o suporte técnico para a internalização de políticas finalísticas no ecossistema de transferências da União, substituindo a análise documental analítica por uma sistemática de adesão rápida.</p>
              
              <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 1.5rem; margin-top: 1.5rem;">
                <div style="background: rgba(0, 242, 254, 0.03); border: 1px solid rgba(0, 242, 254, 0.15); padding: 1.25rem; border-radius: 8px;">
                  <h4 style="color: var(--secondary); margin-bottom: 0.75rem; font-family: var(--font-display); font-size: 1rem; display: flex; align-items: center; gap: 0.5rem;">Papel do MGI</h4>
                  <p style="font-size: 0.85rem; line-height: 1.5; color: var(--text-secondary); margin: 0;">Gestor central da Plataforma TRANSFEREGOV.BR, responsável por fornecer a infraestrutura sistêmica e garantir o marco regulatório e governança.</p>
                </div>
                
                <div style="background: rgba(16, 185, 129, 0.03); border: 1px solid rgba(16, 185, 129, 0.15); padding: 1.25rem; border-radius: 8px;">
                  <h4 style="color: #10b981; margin-bottom: 0.75rem; font-family: var(--font-display); font-size: 1rem; display: flex; align-items: center; gap: 0.5rem;">Papel dos Órgãos Parceiros</h4>
                  <p style="font-size: 0.85rem; line-height: 1.5; color: var(--text-secondary); margin: 0;">Detêm total autonomia na gestão dos recursos de seus fundos, na definição das regras de negócio de suas chamadas públicas e na validação das metas físicas pactuadas.</p>
                </div>
              </div>
            </div>
          </div>

          <!-- Cap 3 -->
          <div class="report-chapter">
            <button class="report-chapter-header" type="button" aria-expanded="false">⚙️ 3. Arquitetura Técnica do TRANSFEREGOV.BR <span>▸</span></button>
            <div class="report-chapter-content">
              <p style="font-size: 0.95rem; line-height: 1.6; margin-bottom: 1.5rem; color: var(--text-secondary);">A robustez do TRANSFEREGOV.BR garante a integridade dos recursos federais. A arquitetura sistêmica é subdividida em três ambientes principais:</p>
              
              <div style="display: grid; grid-template-columns: repeat(3, 1fr); gap: 1rem; margin-bottom: 1.5rem;">
                <div style="background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.05); padding: 1.25rem 1rem; border-radius: 8px; text-align: center;">
                  <span style="background: rgba(0, 242, 254, 0.1); color: var(--secondary); padding: 0.25rem 0.5rem; border-radius: 4px; font-size: 0.75rem; font-weight: bold; text-transform: uppercase;">1. Treinamento</span>
                  <p style="font-size: 0.8rem; color: var(--text-secondary); margin-top: 0.75rem; line-height: 1.4;">Simulador (sandbox) isolado destinado ao aprendizado das equipes técnicas, sem impacto orçamentário real.</p>
                </div>
                <div style="background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.05); padding: 1.25rem 1rem; border-radius: 8px; text-align: center;">
                  <span style="background: rgba(255, 193, 7, 0.1); color: #ffc107; padding: 0.25rem 0.5rem; border-radius: 4px; font-size: 0.75rem; font-weight: bold; text-transform: uppercase;">2. Homologação</span>
                  <p style="font-size: 0.8rem; color: var(--text-secondary); margin-top: 0.75rem; line-height: 1.4;">Validações finais de layouts, parametrizações e testes de integração de APIs com agentes bancários.</p>
                </div>
                <div style="background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.05); padding: 1.25rem 1rem; border-radius: 8px; text-align: center;">
                  <span style="background: rgba(220, 53, 69, 0.1); color: #dc3545; padding: 0.25rem 0.5rem; border-radius: 4px; font-size: 0.75rem; font-weight: bold; text-transform: uppercase;">3. Produção</span>
                  <p style="font-size: 0.8rem; color: var(--text-secondary); margin-top: 0.75rem; line-height: 1.4;">Execução oficial das parcerias, com empenho no SIAFI, transações financeiras reais e monitoramento legal.</p>
                </div>
              </div>
              
              <div style="background: rgba(255, 255, 255, 0.03); border: 1px dashed var(--border-color); padding: 1rem; border-radius: 6px; font-size: 0.85rem; display: flex; align-items: center; gap: 0.75rem; color: var(--text-secondary);">
                <span style="font-size: 1.25rem;" aria-hidden="true">🔧</span>
                <span><strong>Solução Técnica SIORG:</strong> Para contornar eventuais ausências de um código de Unidade Gestora próprio, a plataforma utiliza a vinculação via código de SIORG (ex: código do Departamento Gestor de Recursos - 267384) e perfis específicos para servidores.</span>
              </div>
            </div>
          </div>

          <!-- Cap 4 -->
          <div class="report-chapter">
            <button class="report-chapter-header" type="button" aria-expanded="false">🔧 4. Parametrização por Política Pública <span>▸</span></button>
            <div class="report-chapter-content">
              <p style="font-size: 0.95rem; line-height: 1.6; margin-bottom: 1.5rem; color: var(--text-secondary);">O Módulo de Gestão de Parcerias do TRANSFEREGOV.BR se adapta às particularidades de diferentes políticas públicas por meio de configurações no sistema:</p>
              
              <div style="display: flex; flex-direction: column; gap: 1rem;">
                <div style="display: flex; gap: 1rem; align-items: flex-start; background: rgba(0,0,0,0.15); padding: 1rem; border-radius: 8px;">
                  <div style="background: var(--secondary); color: var(--bg-deep); width: 28px; height: 28px; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-weight: bold; flex-shrink: 0;" aria-hidden="true">1</div>
                  <div>
                    <h5 style="color: #fff; margin-bottom: 0.25rem; font-size: 0.95rem; font-family: var(--font-display);">Requisitos de Habilitação</h5>
                    <p style="font-size: 0.85rem; color: var(--text-secondary); margin: 0; line-height: 1.4;">Definição automatizada de critérios técnicos e fiscais que os proponentes devem cumprir eletronicamente.</p>
                  </div>
                </div>
                
                <div style="display: flex; gap: 1rem; align-items: flex-start; background: rgba(0,0,0,0.15); padding: 1rem; border-radius: 8px;">
                  <div style="background: var(--secondary); color: var(--bg-deep); width: 28px; height: 28px; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-weight: bold; flex-shrink: 0;" aria-hidden="true">2</div>
                  <div>
                    <h5 style="color: #fff; margin-bottom: 0.25rem; font-size: 0.95rem; font-family: var(--font-display);">Anexos Obrigatórios</h5>
                    <p style="font-size: 0.85rem; color: var(--text-secondary); margin: 0; line-height: 1.4;">Possibilidade de restringir a adesão à inserção de planos setoriais específicos de trabalho (ex: planos de contingência, declarações locais).</p>
                  </div>
                </div>
                
                <div style="display: flex; gap: 1rem; align-items: flex-start; background: rgba(0,0,0,0.15); padding: 1rem; border-radius: 8px;">
                  <div style="background: var(--secondary); color: var(--bg-deep); width: 28px; height: 28px; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-weight: bold; flex-shrink: 0;" aria-hidden="true">3</div>
                  <div>
                    <h5 style="color: #fff; margin-bottom: 0.25rem; font-size: 0.95rem; font-family: var(--font-display);">Metas Padronizadas</h5>
                    <p style="font-size: 0.85rem; color: var(--text-secondary); margin: 0; line-height: 1.4;">Catálogo de entregas físicas predefinidas para a rápida adesão dos estados e municípios, eliminando propostas subjetivas.</p>
                  </div>
                </div>
              </div>
            </div>
          </div>

          <!-- Cap 5 -->
          <div class="report-chapter">
            <button class="report-chapter-header" type="button" aria-expanded="false">💳 5. Ordem de Pagamento de Parcerias (OPP) <span>▸</span></button>
            <div class="report-chapter-content">
              <div class="cartilha-grid" style="display: grid; grid-template-columns: 1.2fr 0.8fr; gap: 2rem; align-items: start; margin-top: 1rem;">
                <div>
                  <p style="font-size: 0.95rem; line-height: 1.6; margin-bottom: 1.5rem; color: var(--text-secondary);">A OPP é o instrumento financeiro definitivo de automatização bancária das parcerias. Ela assegura e agiliza a liberação dos recursos federais:</p>
                  
                  <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 1rem; margin-bottom: 1.5rem;">
                    <div style="background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.05); padding: 1rem; border-radius: 6px;">
                      <h5 style="color: var(--secondary); margin-bottom: 0.25rem; font-size: 0.9rem; font-family: var(--font-display);">⚡ Tempo Real via API</h5>
                      <p style="font-size: 0.75rem; color: var(--text-secondary); line-height: 1.4; margin: 0;">Conexão direta com bancos públicos oficiais, eliminando o processamento em lote (batch).</p>
                    </div>
                    <div style="background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.05); padding: 1rem; border-radius: 6px;">
                      <h5 style="color: var(--secondary); margin-bottom: 0.25rem; font-size: 0.9rem; font-family: var(--font-display);">🛡️ Dupla Assinatura</h5>
                      <p style="font-size: 0.75rem; color: var(--text-secondary); line-height: 1.4; margin: 0;">Exigência de dupla autorização eletrônica (Operador e Gestor) com certificação Gov.br Ouro.</p>
                    </div>
                    <div style="background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.05); padding: 1rem; border-radius: 6px;">
                      <h5 style="color: var(--secondary); margin-bottom: 0.25rem; font-size: 0.9rem; font-family: var(--font-display);">📱 Pix e QR Code</h5>
                      <p style="font-size: 0.75rem; color: var(--text-secondary); line-height: 1.4; margin: 0;">Pagamento direto a fornecedores e prestadores na ponta, garantindo conciliação imediata.</p>
                    </div>
                    <div style="background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.05); padding: 1rem; border-radius: 6px;">
                      <h5 style="color: var(--secondary); margin-bottom: 0.25rem; font-size: 0.9rem; font-family: var(--font-display);">⚙️ Controle Ativo</h5>
                      <p style="font-size: 0.75rem; color: var(--text-secondary); line-height: 1.4; margin: 0;">Prerrogativa do órgão federal de comandar o bloqueio e devolução de saldos não utilizados.</p>
                    </div>
                  </div>
                  
                  <div style="background: rgba(255, 193, 7, 0.05); border-left: 3px solid #ffc107; padding: 1rem; border-radius: 4px; font-size: 0.85rem; color: var(--text-secondary); display: flex; flex-direction: column; gap: 0.75rem;">
                    <div><strong>Ficha Técnica e Prazos:</strong> Abertura de conta via API em 3 a 5 segundos; Janela de autorizações das 07h às 19h30; transações via Pix sem teto de valor. A migração compulsória para todas as 63 mil contas ativas no país deve ocorrer até dezembro de 2026.</div>
                    <div style="border-top: 1px dashed rgba(255,193,7,0.2); padding-top: 0.5rem;"><strong>Consulta Pública:</strong> A consulta à OPP no Gestão de Parcerias está acessível na navegação em acesso livre em: <a href="https://parcerias.transferegov.sistema.gov.br/gestaofinanceira/opp/consulta" target="_blank" rel="noopener noreferrer" style="color: var(--secondary); text-decoration: underline;">https://parcerias.transferegov.sistema.gov.br/gestaofinanceira/opp/consulta</a>. Exemplo de Parceria com OPP em produção (Piloto): <strong>2024-00000623</strong>.</div>
                  </div>
                </div>
                
                <div style="text-align: center; background: rgba(0,0,0,0.2); padding: 1rem; border-radius: 8px; border: 1px solid var(--border-color);">
                  <span style="font-size: 0.8rem; color: var(--text-muted); display: block; margin-bottom: 0.5rem; font-weight: bold; text-transform: uppercase; letter-spacing: 1px;">Fluxo Operacional OPP</span>
                  <a href="materiais_complementares/OPP TREINAMENTO/Fluxo_da_Ordem_de_Pagamento.png" target="_blank" rel="noopener noreferrer">
                    <img src="materiais_complementares/OPP TREINAMENTO/Fluxo_da_Ordem_de_Pagamento.png" alt="Diagrama do Fluxo da Ordem de Pagamento de Parcerias" style="width: 100%; border-radius: 4px; box-shadow: 0 4px 12px rgba(0,0,0,0.4); margin-bottom: 0.5rem; transition: transform 0.2s;" onmouseover="this.style.transform='scale(1.02)'" onmouseout="this.style.transform='scale(1)'">
                  </a>
                  <span style="font-size: 0.75rem; color: var(--text-muted); display: block;">Clique para abrir o fluxograma bancário</span>
                </div>
              </div>
            </div>
          </div>

          <!-- Cap 6 -->
          <div class="report-chapter">
            <button class="report-chapter-header" type="button" aria-expanded="false">🗺️ 6. Jornada da Transferência <span>▸</span></button>
            <div class="report-chapter-content">
              <div class="cartilha-grid" style="display: grid; grid-template-columns: 1.2fr 0.8fr; gap: 2rem; align-items: start; margin-top: 1rem;">
                <div>
                  <p style="font-size: 0.95rem; line-height: 1.6; margin-bottom: 1.5rem; color: var(--text-secondary);">A jornada de repasses simplificados substitui os planos de trabalho densos por um fluxo lógico em 5 etapas no sistema:</p>
                  
                  <div style="display: flex; flex-direction: column; gap: 0.75rem;">
                    <div style="display: flex; align-items: center; gap: 1rem; background: rgba(255,255,255,0.02); padding: 0.75rem 1rem; border-radius: 6px; border-left: 3px solid var(--primary); font-size: 0.85rem; color: var(--text-secondary);">
                      <strong style="font-size: 1.2rem; color: var(--primary);">01</strong>
                      <div>
                        <strong>Convocatória:</strong> O órgão concedente cadastra a chamada pública, metas físicas padronizadas e prazos no sistema.
                      </div>
                    </div>
                    <div style="display: flex; align-items: center; gap: 1rem; background: rgba(255,255,255,0.02); padding: 0.75rem 1rem; border-radius: 6px; border-left: 3px solid var(--primary); font-size: 0.85rem; color: var(--text-secondary);">
                      <strong style="font-size: 1.2rem; color: var(--primary);">02</strong>
                      <div>
                        <strong>Manifestação de Interesse:</strong> O ente subnacional adere digitalmente à chamada, preenchendo as informações e anexando os documentos.
                      </div>
                    </div>
                    <div style="display: flex; align-items: center; gap: 1rem; background: rgba(255,255,255,0.02); padding: 0.75rem 1rem; border-radius: 6px; border-left: 3px solid var(--primary); font-size: 0.85rem; color: var(--text-secondary);">
                      <strong style="font-size: 1.2rem; color: var(--primary);">03</strong>
                      <div>
                        <strong>Validação Técnica:</strong> O concedente realiza a análise da habilitação eletrônica e emite o parecer de aprovação.
                      </div>
                    </div>
                    <div style="display: flex; align-items: center; gap: 1rem; background: rgba(255,255,255,0.02); padding: 0.75rem 1rem; border-radius: 6px; border-left: 3px solid var(--primary); font-size: 0.85rem; color: var(--text-secondary);">
                      <strong style="font-size: 1.2rem; color: var(--primary);">04</strong>
                      <div>
                        <strong>Empenho e Desembolso:</strong> O sistema realiza o empenho automático e o comando de abertura de conta bancária vinculada.
                      </div>
                    </div>
                    <div style="display: flex; align-items: center; gap: 1rem; background: rgba(255,255,255,0.02); padding: 0.75rem 1rem; border-radius: 6px; border-left: 3px solid var(--primary); font-size: 0.85rem; color: var(--text-secondary);">
                      <strong style="font-size: 1.2rem; color: var(--primary);">05</strong>
                      <div>
                        <strong>Execução e Monitoramento:</strong> O ente executa as metas físicas e apresenta o Relatório de Execução Simplificado focado nas entregas.
                      </div>
                    </div>
                  </div>
                </div>
                
                <div style="text-align: center; background: rgba(0,0,0,0.2); padding: 1rem; border-radius: 8px; border: 1px solid var(--border-color);">
                  <span style="font-size: 0.8rem; color: var(--text-muted); display: block; margin-bottom: 0.5rem; font-weight: bold; text-transform: uppercase; letter-spacing: 1px;">Mapa Mental do Fluxo</span>
                  <a href="materiais_complementares/slides/Jornada_da_Transferência_Simplificada (1).png" target="_blank" rel="noopener noreferrer">
                    <img src="materiais_complementares/slides/Jornada_da_Transferência_Simplificada (1).png" alt="Esquema da Jornada da Transferência Simplificada" style="width: 100%; border-radius: 4px; box-shadow: 0 4px 12px rgba(0,0,0,0.4); margin-bottom: 0.5rem; transition: transform 0.2s;" onmouseover="this.style.transform='scale(1.02)'" onmouseout="this.style.transform='scale(1)'">
                  </a>
                  <span style="font-size: 0.75rem; color: var(--text-muted); display: block;">Clique para abrir a jornada simplificada</span>
                </div>
              </div>
            </div>
          </div>

          <!-- Cap 7 -->
          <div class="report-chapter">
            <button class="report-chapter-header" type="button" aria-expanded="false">❓ 7. Dúvidas Operacionais (FAQ) <span>▸</span></button>
            <div class="report-chapter-content">
              <div style="display: flex; flex-direction: column; gap: 1.25rem;">
                <div style="background: rgba(255, 255, 255, 0.02); padding: 1.25rem; border-radius: 8px; border: 1px solid var(--border-color);">
                  <p style="color: #fff; font-weight: bold; margin-bottom: 0.5rem; display: flex; align-items: center; gap: 0.5rem; font-family: var(--font-display);">❓ Como funciona a adesão no TRANSFEREGOV.BR?</p>
                  <p style="font-size: 0.85rem; color: var(--text-secondary); line-height: 1.5; margin: 0; padding-left: 1.5rem;">O processo é eletrônico e integrado. O proponente localiza o edital e manifesta interesse aderindo diretamente no TRANSFEREGOV.BR, assinando digitalmente o termo.</p>
                </div>
                
                <div style="background: rgba(255, 255, 255, 0.02); padding: 1.25rem; border-radius: 8px; border: 1px solid var(--border-color);">
                  <p style="color: #fff; font-weight: bold; margin-bottom: 0.5rem; display: flex; align-items: center; gap: 0.5rem; font-family: var(--font-display);">❓ O que significa parametrizar o edital?</p>
                  <p style="font-size: 0.85rem; color: var(--text-secondary); line-height: 1.5; margin: 0; padding-left: 1.5rem;">É a predefinição pelo órgão concedente das regras de negócio, metas padronizadas e documentos exigidos, adaptando o sistema ao seu setor.</p>
                </div>
                
                <div style="background: rgba(255, 255, 255, 0.02); padding: 1.25rem; border-radius: 8px; border: 1px solid var(--border-color);">
                  <p style="color: #fff; font-weight: bold; margin-bottom: 0.5rem; display: flex; align-items: center; gap: 0.5rem; font-family: var(--font-display);">❓ Qual a importância da Sandbox de Treinamento?</p>
                  <p style="font-size: 0.85rem; color: var(--text-secondary); line-height: 1.5; margin: 0; padding-left: 1.5rem;">Funciona como um simulador (sandbox) isolado, permitindo testes práticos de empenho e execução de forma prática e sem risco fiscal real.</p>
                </div>
                
                <div style="background: rgba(255, 255, 255, 0.02); padding: 1.25rem; border-radius: 8px; border: 1px solid var(--border-color);">
                  <p style="color: #fff; font-weight: bold; margin-bottom: 0.5rem; display: flex; align-items: center; gap: 0.5rem; font-family: var(--font-display);">❓ Como ocorre a prestação de contas?</p>
                  <p style="font-size: 0.85rem; color: var(--text-secondary); line-height: 1.5; margin: 0; padding-left: 1.5rem;">O foco é direcionado para a comprovação física do alcance do objeto e das metas cadastradas, simplificando a prestação documental ex-post.</p>
                </div>
              </div>
            </div>
          </div>

          <!-- Cap 8 -->
          <div class="report-chapter">
            <button class="report-chapter-header" type="button" aria-expanded="false">⚠️ 8. Pontos de Atenção e Riscos <span>▸</span></button>
            <div class="report-chapter-content">
              <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 1.5rem; margin-top: 0.5rem;">
                <div style="background: rgba(220, 53, 69, 0.03); border: 1px solid rgba(220, 53, 69, 0.15); border-left: 4px solid #dc3545; padding: 1.25rem; border-radius: 8px;">
                  <h4 style="color: #dc3545; margin-bottom: 0.5rem; font-family: var(--font-display); font-size: 1rem; display: flex; align-items: center; gap: 0.5rem;">
                    <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" aria-hidden="true"><path d="M10.29 3.86L1.82 18a2 2 0 0 0 1.71 3h16.94a2 2 0 0 0 1.71-3L13.71 3.86a2 2 0 0 0-3.42 0z"></path><line x1="12" y1="9" x2="12" y2="13"></line><line x1="12" y1="17" x2="12.01" y2="17"></line></svg>
                    Sintoma do Convênio
                  </h4>
                  <p style="font-size: 0.85rem; line-height: 1.5; color: var(--text-secondary); margin: 0;">Dificuldade cultural de técnicos em desapegar de controles ex-ante densos e Notas Fiscais unitárias ex-ante, o que invalida a celeridade proposta pela Transferência Simplificada.</p>
                </div>
                
                <div style="background: rgba(255, 193, 7, 0.03); border: 1px solid rgba(255, 193, 7, 0.15); border-left: 4px solid #ffc107; padding: 1.25rem; border-radius: 8px;">
                  <h4 style="color: #ffc107; margin-bottom: 0.5rem; font-family: var(--font-display); font-size: 1rem; display: flex; align-items: center; gap: 0.5rem;">
                    <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" aria-hidden="true"><circle cx="12" cy="12" r="10"></circle><line x1="12" y1="8" x2="12" y2="12"></line><line x1="12" y1="16" x2="12.01" y2="16"></line></svg>
                    Prazos do Defeso Eleitoral
                  </h4>
                  <p style="font-size: 0.85rem; line-height: 1.5; color: var(--text-secondary); margin: 0;">Observância estrita do calendário eleitoral ordinário. Importante notar que apenas programas de Emergência ou Calamidade Pública formalmente decretada por portaria ministerial ativa possuem exceções de continuidade.</p>
                </div>
              </div>
            </div>
          </div>

          <!-- Cap 9 -->
          <div class="report-chapter">
            <button class="report-chapter-header" type="button" aria-expanded="false">✅ 9. Boas Práticas Operacionais <span>▸</span></button>
            <div class="report-chapter-content">
              <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 1.5rem; margin-top: 0.5rem;">
                <div style="background: rgba(16, 185, 129, 0.03); border: 1px solid rgba(16, 185, 129, 0.15); padding: 1.25rem; border-radius: 8px;">
                  <h4 style="color: #10b981; margin-bottom: 0.5rem; font-family: var(--font-display); font-size: 1rem; display: flex; align-items: center; gap: 0.5rem;">
                    <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" aria-hidden="true"><polyline points="20 6 9 17 4 12"></polyline></svg>
                    Treinamento Prévio em Sandbox
                  </h4>
                  <p style="font-size: 0.85rem; line-height: 1.5; color: var(--text-secondary); margin: 0;">Utilização intensiva do ambiente de treinamento livre para simular 90% das telas sem necessidade imediata de certificado digital.</p>
                </div>
                
                <div style="background: rgba(16, 185, 129, 0.03); border: 1px solid rgba(16, 185, 129, 0.15); padding: 1.25rem; border-radius: 8px;">
                  <h4 style="color: #10b981; margin-bottom: 0.5rem; font-family: var(--font-display); font-size: 1rem; display: flex; align-items: center; gap: 0.5rem;">
                    <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" aria-hidden="true"><polyline points="20 6 9 17 4 12"></polyline></svg>
                    Diálogo Técnico Permanente
                  </h4>
                  <p style="font-size: 0.85rem; line-height: 1.5; color: var(--text-secondary); margin: 0;">Abertura de chats técnicos ou fóruns integrados para esclarecer parametrizações de metas com o órgão concedente antes de abrir o edital.</p>
                </div>
              </div>
            </div>
          </div>"""

    import subprocess
    # Restore index.html to clean git baseline before patching to make it idempotent
    subprocess.run(["git", "checkout", "--", index_path], capture_output=True)

    # We will locate the target section in index.html and overwrite it
    with open(index_path, "r", encoding="utf-8") as f:
        html_content = f.read()

    # Find start and end indices of the chapters block
    start_pos = html_content.find("          <!-- Cap 1 -->")
    end_target = "          <!-- Cap 9 -->"
    cap9_pos = html_content.find(end_target)
    
    if start_pos != -1 and cap9_pos != -1:
        cap9_end_pos = html_content.find("          </div>\n        </div>", cap9_pos)
        offset = len("          </div>\n")
        if cap9_end_pos == -1:
            cap9_end_pos = html_content.find("          </div>        </div>", cap9_pos)
            offset = len("          </div>")
            
        if cap9_end_pos != -1:
            cap9_end_pos += offset # Include the cap 9 closing div tag
            
            patched_content = html_content[:start_pos] + new_chapters_html + html_content[cap9_end_pos:]
            with open(index_path, "w", encoding="utf-8") as f:
                f.write(patched_content)
            print("index.html modal chapters patched successfully!")
        else:
            print("Error: Could not locate the closing container of Cap 9.")
    else:
        print("Error: Could not find start or end tags of chapters inside index.html.")

if __name__ == '__main__':
    generate_docx()
    generate_html_files()
    patch_index_modal()
